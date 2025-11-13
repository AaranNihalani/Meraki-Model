import os, re, glob

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False

def _read_docx_zip_text(path):
    try:
        from zipfile import ZipFile
        import xml.etree.ElementTree as ET
        with ZipFile(path) as z:
            xml = z.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        texts = [el.text for el in root.findall(".//w:t", ns) if el.text]
        txt = " ".join(texts)
        return re.sub(r"\s+", " ", txt.strip())
    except Exception:
        return ""

def read_docx(path):
    if not DOCX_AVAILABLE:
        print("⚠️ python-docx not installed; skipping DOCX extraction")
        return ""
    doc = Document(path)
    parts = []
    parts.extend([p.text for p in doc.paragraphs if p.text and p.text.strip()])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)
    for section in getattr(doc, "sections", []):
        for hf in [section.header, section.footer]:
            for p in hf.paragraphs:
                if p.text and p.text.strip():
                    parts.append(p.text)
    text = re.sub(r"\s+", " ", "\n".join(parts).strip())
    if not text:
        text = _read_docx_zip_text(path)
    return text

def read_pdf(path):
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return re.sub(r"\s+", " ", text.strip())
    except Exception:
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(path)
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
            return re.sub(r"\s+", " ", text.strip())
        except Exception:
            try:
                from pdfminer.high_level import extract_text
                text = extract_text(path) or ""
                return re.sub(r"\s+", " ", text.strip())
            except Exception:
                return ""

def main():
    in_dir = "./data/domain_docs"
    out_dir = "./data/domain_docs_labeled"
    os.makedirs(out_dir, exist_ok=True)

    files = glob.glob(os.path.join(in_dir, "*.docx")) + glob.glob(os.path.join(in_dir, "*.pdf"))
    if not files:
        print("⚠️ No files found in ./data/domain_docs")
        return

    count = 0
    for path in files:
        base = os.path.splitext(os.path.basename(path))[0]
        out_path = os.path.join(out_dir, base + ".txt")
        text = ""
        if path.lower().endswith(".docx"):
            text = read_docx(path)
        elif path.lower().endswith(".pdf"):
            text = read_pdf(path)
        if text:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(text)
            count += 1
            print(f"💾 Saved: {out_path}")
        else:
            print(f"⚠️ Skipped (no text): {path}")

    print(f"✅ Processed {count} files → {out_dir}")

if __name__ == "__main__":
    main()